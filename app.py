# from flask import Flask, render_template, request, jsonify
# from PyPDF2 import PdfReader
# from langchain.embeddings.openai import OpenAIEmbeddings
# from langchain.text_splitter import CharacterTextSplitter
# from langchain.vectorstores import FAISS
# from langchain.chains.question_answering import load_qa_chain
# from langchain.llms import OpenAI
# from pdf2image import convert_from_path
# import pytesseract
# import os
# import base64
# from gtts import gTTS
# from io import BytesIO
# import requests

# app = Flask(__name__)

# # Set up environment variables
# os.environ["OPENAI_API_KEY"] = "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"
# GOOGLE_API_KEY = "AIzaSyBUfRJiBynN-YN0JOoDFAXmgTBCXHp9es4"

# # Configure pytesseract to use the Tesseract executable
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# # Load OpenAI embeddings and set up document search
# embeddings = OpenAIEmbeddings()
# document_search = None  # Define globally

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/ask', methods=['POST'])
# def ask_question():
#     if request.method == 'POST':
#         pdf_file = request.files.get('pdf')
#         query = request.form.get('question')
        
#         if not pdf_file or not query:
#             return jsonify({'error': 'PDF file and question are required'}), 400

#         # Read text from PDF
#         raw_text = ''
#         try:
#             pdf_reader = PdfReader(pdf_file)
#             for page in pdf_reader.pages:
#                 content = page.extract_text()
#                 if content:
#                     raw_text += content

#             # If no text was extracted, use OCR
#             if not raw_text.strip():
#                 images = convert_from_path(pdf_file)
#                 for image in images:
#                     raw_text += pytesseract.image_to_string(image)
#         except Exception as e:
#             return jsonify({'error': 'Failed to read PDF file', 'details': str(e)}), 500

#         if not raw_text:
#             return jsonify({'error': 'No text found in PDF'}), 400

#         # Split text into chunks for processing
#         text_splitter = CharacterTextSplitter(
#             separator="\n",
#             chunk_size=800,
#             chunk_overlap=200,
#             length_function=len,
#         )
#         texts = text_splitter.split_text(raw_text)

#         if not texts:
#             return jsonify({'error': 'Text splitting resulted in no chunks'}), 400

#         # Perform document search
#         global document_search
#         try:
#             document_search = FAISS.from_texts(texts, embeddings)
#         except IndexError as e:
#             return jsonify({'error': 'Error creating embeddings', 'details': str(e)}), 500

#         # Load question-answering chain
#         chain = load_qa_chain(OpenAI(), chain_type="stuff")

#         # Search for similar documents and get the answer
#         docs = document_search.similarity_search(query)
#         answer = chain.run(input_documents=docs, question=query)

#         return jsonify({'answer': answer})

# @app.route('/tts', methods=['POST'])
# def tts():
#     data = request.get_json()
#     text = data.get('text', '')
#     if not text:
#         return jsonify({'error': 'No text provided for TTS'}), 400

#     try:
#         tts = gTTS(text, lang='en')
#         mp3_fp = BytesIO()
#         tts.write_to_fp(mp3_fp)
#         mp3_fp.seek(0)
#         audio_base64 = base64.b64encode(mp3_fp.read()).decode('utf-8')
#     except Exception as e:
#         return jsonify({'error': 'Failed to generate audio', 'details': str(e)}), 500

#     return jsonify({'audio': audio_base64})

# @app.route('/translate', methods=['POST'])
# def translate_text():
#     data = request.get_json()
#     text = data.get('text', '')
#     target_language = data.get('language', '')

#     if not text or not target_language:
#         return jsonify({'error': 'Text and target language are required'}), 400

#     try:
#         response = requests.post(
#             'https://translation.googleapis.com/language/translate/v2',
#             params={
#                 'q': text,
#                 'target': target_language,
#                 'key': GOOGLE_API_KEY
#             }
#         )
#         response.raise_for_status()
#         translated_text = response.json()['data']['translations'][0]['translatedText']
#     except requests.exceptions.HTTPError as http_err:
#         return jsonify({'error': 'HTTP error occurred', 'details': str(http_err)}), 500
#     except Exception as e:
#         return jsonify({'error': 'Failed to translate text', 'details': str(e)}), 500

#     return jsonify({'translated_text': translated_text})

# @app.route('/download', methods=['POST'])
# def download():
#     data = request.get_json()
#     text = data.get('text', '')
#     if not text:
#         return jsonify({'error': 'No text provided for download'}), 400

#     try:
#         from docx import Document
#         doc = Document()
#         doc.add_paragraph(text)
#         file_stream = BytesIO()
#         doc.save(file_stream)
#         file_stream.seek(0)
#         file_base64 = base64.b64encode(file_stream.read()).decode('utf-8')
#     except Exception as e:
#         return jsonify({'error': 'Failed to generate document', 'details': str(e)}), 500

#     return jsonify({'file': file_base64})

# if __name__ == '__main__':
#     app.run(debug=True)
from flask import Flask, render_template, request, jsonify
from PyPDF2 import PdfReader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from pdf2image import convert_from_path
import pytesseract
import os
import base64
from gtts import gTTS
from io import BytesIO
import requests
import tempfile

app = Flask(__name__)

# Set up environment variables
os.environ["OPENAI_API_KEY"] = "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"
GOOGLE_API_KEY = "AIzaSyBUfRJiBynN-YN0JOoDFAXmgTBCXHp9es4"

# Configure pytesseract to use the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Load OpenAI embeddings and set up document search
embeddings = OpenAIEmbeddings()
document_search = None  # Define globally

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/ask', methods=['POST'])
def ask_question():
    if request.method == 'POST':
        pdf_file = request.files.get('pdf')
        query = request.form.get('question')
        
        if not pdf_file or not query:
            return jsonify({'error': 'PDF file and question are required'}), 400

        # Save the uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            pdf_path = tmp.name
            pdf_file.save(pdf_path)

        # Read text from PDF
        raw_text = ''
        try:
            pdf_reader = PdfReader(pdf_path)
            for page in pdf_reader.pages:
                content = page.extract_text()
                if content:
                    raw_text += content

            # If no text was extracted, use OCR
            if not raw_text.strip():
                images = convert_from_path(pdf_path)
                for image in images:
                    raw_text += pytesseract.image_to_string(image)
        except Exception as e:
            return jsonify({'error': 'Failed to read PDF file', 'details': str(e)}), 500
        finally:
            os.remove(pdf_path)  # Clean up the temporary file

        if not raw_text:
            return jsonify({'error': 'No text found in PDF'}), 400

        # Split text into chunks for processing
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=800,
            chunk_overlap=200,
            length_function=len,
        )
        texts = text_splitter.split_text(raw_text)

        if not texts:
            return jsonify({'error': 'Text splitting resulted in no chunks'}), 400

        # Perform document search
        global document_search
        try:
            document_search = FAISS.from_texts(texts, embeddings)
        except IndexError as e:
            return jsonify({'error': 'Error creating embeddings', 'details': str(e)}), 500

        # Load question-answering chain
        chain = load_qa_chain(OpenAI(), chain_type="stuff")

        # Search for similar documents and get the answer
        docs = document_search.similarity_search(query)
        answer = chain.run(input_documents=docs, question=query)

        return jsonify({'answer': answer})

@app.route('/tts', methods=['POST'])
def tts():
    data = request.get_json()
    text = data.get('text', '')
    if not text:
        return jsonify({'error': 'No text provided for TTS'}), 400

    try:
        tts = gTTS(text, lang='en')
        mp3_fp = BytesIO()
        tts.write_to_fp(mp3_fp)
        mp3_fp.seek(0)
        audio_base64 = base64.b64encode(mp3_fp.read()).decode('utf-8')
    except Exception as e:
        return jsonify({'error': 'Failed to generate audio', 'details': str(e)}), 500

    return jsonify({'audio': audio_base64})

@app.route('/translate', methods=['POST'])
def translate_text():
    data = request.get_json()
    text = data.get('text', '')
    target_language = data.get('language', '')

    if not text or not target_language:
        return jsonify({'error': 'Text and target language are required'}), 400

    try:
        response = requests.post(
            'https://translation.googleapis.com/language/translate/v2',
            params={
                'q': text,
                'target': target_language,
                'key': GOOGLE_API_KEY
            }
        )
        response.raise_for_status()
        translated_text = response.json()['data']['translations'][0]['translatedText']
    except requests.exceptions.HTTPError as http_err:
        return jsonify({'error': 'HTTP error occurred', 'details': str(http_err)}), 500
    except Exception as e:
        return jsonify({'error': 'Failed to translate text', 'details': str(e)}), 500

    return jsonify({'translated_text': translated_text})

@app.route('/download', methods=['POST'])
def download():
    data = request.get_json()
    text = data.get('text', '')
    if not text:
        return jsonify({'error': 'No text provided for download'}), 400

    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_base64 = base64.b64encode(file_stream.read()).decode('utf-8')
    except Exception as e:
        return jsonify({'error': 'Failed to generate document', 'details': str(e)}), 500

    return jsonify({'file': file_base64})

if __name__ == '__main__':
    app.run(debug=True)
